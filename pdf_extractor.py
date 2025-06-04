import json
import base64
import os
import re
from datetime import datetime
from typing import Dict, List, Any, Optional
from dataclasses import dataclass, asdict
import fitz  # PyMuPDF
from PIL import Image
import pytesseract
import cv2
import numpy as np
from pathlib import Path
import requests
import io
import zipfile
import tempfile

# Optional imports for different document types
try:
    from docx import Document
    from docx.table import Table as DocxTable
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False

try:
    import camelot
    CAMELOT_AVAILABLE = True
except ImportError:
    CAMELOT_AVAILABLE = False

@dataclass
class ExtractedTable:
    table_id: str
    headers: List[str]
    rows: List[Dict[str, Any]]
    confidence: float = 0.0

@dataclass
class ExtractedImage:
    image_id: str
    description: str
    position: str
    estimated_size: str
    file_path: str
    metadata: Dict[str, Any]

@dataclass
class PageContent:
    page_number: int
    content_type: str
    text_content: Dict[str, Any]
    tables: List[ExtractedTable]
    images: List[ExtractedImage]

@dataclass
class DocumentMetadata:
    filename: str
    total_pages: int
    extraction_timestamp: str
    file_size: str
    file_type: str

@dataclass
class ExtractionConfig:
    """Configuration for document extraction"""
    use_ai_service: bool = False
    ai_api_key: Optional[str] = None
    ai_service_url: Optional[str] = None
    output_dir: str = "extracted_content"
    save_images: bool = True
    ocr_language: str = "eng"
    table_extraction_method: str = "auto"  # auto, pymupdf, camelot, ai
    min_table_rows: int = 2
    min_table_cols: int = 2
    image_dpi: int = 300
    enable_preprocessing: bool = True

class DocumentExtractor:
    def __init__(self, config: Optional[ExtractionConfig] = None):
        self.config = config or ExtractionConfig()
        self.output_dir = Path(self.config.output_dir)
        self.output_dir.mkdir(exist_ok=True)
        self.images_dir = self.output_dir/"images"
        self.images_dir.mkdir(exist_ok=True)
        
        # Validate dependencies
        if not DOCX_AVAILABLE and self.config.use_ai_service:
            print("Warning: python-docx not available. Word document support limited.")
        
        if not CAMELOT_AVAILABLE and self.config.table_extraction_method == "camelot":
            print("Warning: camelot-py not available. Falling back to pymupdf for tables.")
            self.config.table_extraction_method = "pymupdf"
        
    def extract_document(self, file_path: str) -> Dict[str, Any]:
        """Main extraction method that routes to appropriate handler"""
        file_type = self._get_file_type(file_path)
        
        if file_type == 'pdf':
            return self._extract_pdf(file_path)
        elif file_type in ['jpg', 'jpeg', 'png', 'tiff']:
            return self._extract_image(file_path)
        elif file_type in ['docx', 'doc']:
            return self._extract_word(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
    
    def _extract_pdf(self, file_path: str) -> Dict[str, Any]:
        """Extract content from PDF using multiple methods"""
        doc = fitz.open(file_path)
        pages = []
        
        for page_num in range(len(doc)):
            page = doc[page_num]
            
            # Extract text blocks
            text_blocks = page.get_text("dict")
            
            # Extract images
            image_list = page.get_images()
            extracted_images = self._process_page_images(page, image_list, page_num, file_path)
            
            # Extract tables (using layout analysis)
            tables = self._extract_tables_from_page(page, page_num, file_path)
            
            # Analyze page layout and content type
            content_type = self._analyze_content_type(text_blocks, tables, extracted_images)
            
            # Structure text content
            text_content = self._structure_text_content(text_blocks)
            
            page_content = PageContent(
                page_number=page_num + 1,
                content_type=content_type,
                text_content=text_content,
                tables=tables,
                images=extracted_images
            )
            consolidated = self._consolidate_page_content(page_content)
            page_dict = asdict(page_content)
            page_dict["consolidated_content"] = consolidated
            pages.append(page_dict)
        
        # Generate document metadata
        metadata = self._generate_metadata(file_path, len(doc))
        # Create summary
        summary = self._generate_summary_from_dicts(pages)  # Convert back for summary
        
        return {
            "document_metadata": asdict(metadata),
            "pages": pages,
            "summary": summary
        }
    def _generate_summary_from_dicts(self, page_dicts: List[Dict]) -> Dict[str, Any]:
        """Generate document summary from page dictionaries"""
        total_tables = sum(len(page.get('tables', [])) for page in page_dicts)
        total_images = sum(len(page.get('images', [])) for page in page_dicts)
        
        main_topics = []
        key_metrics = []
        
        for page in page_dicts:
            text_content = page.get('text_content', {})
            if text_content.get("title"):
                main_topics.append(text_content["title"])
            
            for table in page.get('tables', []):
                for row in table.get('rows', []):
                    for key, value in row.items():
                        if any(metric_word in str(value).lower() 
                            for metric_word in ["target", "goal", "%", "reduction"]):
                            key_metrics.append(f"{key}: {value}")
        
        return {
            "total_tables": total_tables,
            "total_images": total_images,
            "main_topics": main_topics[:10],
            "key_metrics": key_metrics[:10]
        }
    
    def _extract_tables_from_page(self, page, page_num: int,file_path: str = None) -> List[ExtractedTable]:
        """Extract tables using PDFPlumber first, then Camelot fallback"""
        tables = []
        
        # Method 1: Try PDFPlumber first (most reliable)
        if file_path:
            tables = self._extract_tables_with_pdfplumber(file_path, page_num)
            
        return tables

    def _extract_tables_with_pdfplumber(self, file_path: str, page_num: int) -> List[ExtractedTable]:
        """Extract tables using PDFPlumber"""
        tables = []
        
        try:
            import pdfplumber
            
            with pdfplumber.open(file_path) as pdf:
                if page_num < len(pdf.pages):
                    page = pdf.pages[page_num]
                    
                    # Extract tables
                    page_tables = page.extract_tables()
                    
                    for i, table_data in enumerate(page_tables):
                        if not table_data or len(table_data) < self.config.min_table_rows:
                            continue
                        
                        # Get headers
                        headers = table_data[0] if table_data[0] else []
                        headers = [str(h).strip() if h else f"Column_{j+1}" for j, h in enumerate(headers)]
                        
                        # Process rows
                        rows = []
                        for row_data in table_data[1:]:
                            if row_data and len(row_data) >= self.config.min_table_cols:
                                row_dict = {}
                                for j, cell in enumerate(row_data):
                                    header_key = headers[j] if j < len(headers) else f"Column_{j+1}"
                                    cell_value = str(cell).strip() if cell else ""
                                    row_dict[header_key] = cell_value
                                
                                if any(value.strip() for value in row_dict.values()):
                                    rows.append(row_dict)
                        
                        if rows:
                            tables.append(ExtractedTable(
                                table_id=f"page_{page_num}_pdfplumber_table_{i}",
                                headers=headers,
                                rows=rows,
                                confidence=0.9
                            ))
                            
        except Exception as e:
            print(f"PDFPlumber extraction failed for page {page_num}: {str(e)}")
        
        return tables
    
    def _extract_tables_with_ai(self, page, page_num: int) -> List[ExtractedTable]:
        """Use AI service to extract tables from complex layouts"""
        if not self.config.use_ai_service or not self.config.ai_api_key:
            return []
        
        try:
            # Convert page to image
            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))  # Higher resolution
            img_data = pix.tobytes("png")
            
            # Encode image for API
            img_base64 = base64.b64encode(img_data).decode('utf-8')
            
            # Prepare prompt for AI service
            prompt = """
            Analyze this document page image and extract any tables you find. 
            Return the results in JSON format with this structure:
            {
                "tables": [
                    {
                        "headers": ["col1", "col2", "col3"],
                        "rows": [
                            {"col1": "value1", "col2": "value2", "col3": "value3"},
                            {"col1": "value4", "col2": "value5", "col3": "value6"}
                        ]
                    }
                ]
            }
            If no tables are found, return {"tables": []}.
            """
            
            # Example API call (adjust based on your AI service)
            headers = {
                "Authorization": f"Bearer {self.config.ai_api_key}",
                "Content-Type": "application/json"
            }
            
            payload = {
                "model": "gpt-4-vision-preview",  # or your preferred model
                "messages": [
                    {
                        "role": "user",
                        "content": [
                            {"type": "text", "text": prompt},
                            {"type": "image_url", "image_url": {"url": f"data:image/png;base64,{img_base64}"}}
                        ]
                    }
                ],
                "max_tokens": 2000
            }
            
            if self.config.ai_service_url:
                response = requests.post(self.config.ai_service_url, headers=headers, json=payload)
                
                if response.status_code == 200:
                    result = response.json()
                    content = result.get("choices", [{}])[0].get("message", {}).get("content", "")
                    
                    # Parse JSON response
                    try:
                        tables_data = json.loads(content)
                        extracted_tables = []
                        
                        for i, table_data in enumerate(tables_data.get("tables", [])):
                            extracted_table = ExtractedTable(
                                table_id=f"page_{page_num}_ai_table_{i}",
                                headers=table_data.get("headers", []),
                                rows=table_data.get("rows", []),
                                confidence=0.9
                            )
                            extracted_tables.append(extracted_table)
                        
                        return extracted_tables
                    except json.JSONDecodeError:
                        print(f"Failed to parse AI response for page {page_num}")
                        return []
            
        except Exception as e:
            print(f"AI table extraction failed for page {page_num}: {str(e)}")
            return []
        
        return []
    
    def _process_page_images(self, page, image_list: List, page_num: int, source_file_path: str) -> List[ExtractedImage]:
        """Extract and analyze images from page"""
        extracted_images = []
        
        for img_index, img in enumerate(image_list):
            try:
                # Extract image
                xref = img[0]
                pix = fitz.Pixmap(page.parent, xref)
                
                if pix.n - pix.alpha < 4:  # GRAY or RGB
                    img_data = pix.tobytes("png")
                    doc_name = Path(source_file_path).stem
                    img_path = self.images_dir / f"{doc_name}_page_{page_num}_img_{img_index}.png"                    
                    # Save image (in real implementation)
                if self.config.save_images:
                    try:
                        with open(img_path, "wb") as f:
                            f.write(img_data)
                        print(f"Saved image: {img_path}")
                    except Exception as e:
                        print(f"Failed to save image {img_path}: {str(e)}")

                    
                    description = self._generate_image_description(img_data)
                    
                    extracted_image = ExtractedImage(
                        image_id=f"page_{page_num}_img_{img_index}",
                        description=description,
                        position="embedded",
                        estimated_size=f"{pix.width}x{pix.height}",
                        file_path=str(img_path),
                        metadata={
                            "width": pix.width,
                            "height": pix.height,
                            "colorspace": pix.colorspace.name if pix.colorspace else "unknown",
                            "source_document": os.path.basename(source_file_path),
                            "source_document_full_path": source_file_path,
                            "extraction_timestamp": datetime.now().isoformat(),
                            "page_number": page_num + 1
                        }
                    )
                    extracted_images.append(extracted_image)
                
                pix = None
            except:
                continue
        
        return extracted_images
    
    def _generate_image_description(self, img_data: bytes) -> str:
        """Generate image description using AI or basic analysis"""
        if self.config.use_ai_service and self.config.ai_api_key:
            try:
                # Use AI service to describe image
                img_base64 = base64.b64encode(img_data).decode('utf-8')
                
                prompt = """
                Describe this image briefly and professionally. Focus on:
                1. What type of visual element it is (chart, diagram, photo, icon, etc.)
                2. Main subject or content
                3. Key visual elements or data if applicable
                
                Keep the description concise (1-2 sentences).
                """
                
                headers = {
                    "Authorization": f"Bearer {self.config.ai_api_key}",
                    "Content-Type": "application/json"
                }
                
                payload = {
                    "model": "gpt-4-vision-preview",
                    "messages": [
                        {
                            "role": "user",
                            "content": [
                                {"type": "text", "text": prompt},
                                {"type": "image_url", "image_url": {"url": f"data:image/png;base64,{img_base64}"}}
                            ]
                        }
                    ],
                    "max_tokens": 200
                }
                
                if self.config.ai_service_url:
                    response = requests.post(self.config.ai_service_url, headers=headers, json=payload)
                    if response.status_code == 200:
                        result = response.json()
                        return result.get("choices", [{}])[0].get("message", {}).get("content", "").strip()
            
            except Exception as e:
                print(f"AI image description failed: {str(e)}")
        
        # Fallback to basic image analysis
        try:
            # Convert bytes to numpy array
            img_array = np.frombuffer(img_data, np.uint8)
            img = cv2.imdecode(img_array, cv2.IMREAD_COLOR)
            
            if img is None:
                return "Extracted image"
            
            height, width = img.shape[:2]
            
            # Basic analysis
            if width > height * 2:
                aspect_description = "wide landscape"
            elif height > width * 2:
                aspect_description = "tall portrait"
            else:
                aspect_description = "rectangular"
            
            # Simple color analysis
            avg_color = np.mean(img, axis=(0, 1))
            if np.mean(avg_color) < 50:
                color_description = "dark"
            elif np.mean(avg_color) > 200:
                color_description = "light"
            else:
                color_description = "medium-toned"
            
            # Check for potential chart/graph patterns
            gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
            edges = cv2.Canny(gray, 50, 150)
            contours, _ = cv2.findContours(edges, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
            
            if len(contours) > 10:
                content_type = "diagram or chart"
            elif len(contours) > 5:
                content_type = "structured graphic"
            else:
                content_type = "simple image"
            
            return f"{color_description} {aspect_description} {content_type}"
            
        except Exception as e:
            print(f"Basic image analysis failed: {str(e)}")
            return "Extracted image"
    
    def _structure_text_content(self, text_blocks: Dict) -> Dict[str, Any]:
        """Structure text content by analyzing fonts, positions, etc."""
        content = {
            "title": "",
            "sections": [],
            "main_text": ""
        }
        
        blocks = text_blocks.get("blocks", [])
        current_section = None
        all_text_parts = []  # Collect all text to build coherent content
        
        # First pass: collect all text with metadata
        text_elements = []
        for block in blocks:
            if "lines" in block:
                for line in block["lines"]:
                    line_text = ""
                    max_font_size = 0
                    is_bold = False
                    
                    for span in line["spans"]:
                        text = span["text"].strip()
                        if text:
                            line_text += text + " "
                            max_font_size = max(max_font_size, span["size"])
                            if span["flags"] & 2**4:  # Bold
                                is_bold = True
                    
                    if line_text.strip():
                        text_elements.append({
                            "text": line_text.strip(),
                            "font_size": max_font_size,
                            "is_bold": is_bold,
                            "y_position": line.get("bbox", [0,0,0,0])[1]  # Y coordinate for ordering
                        })
        
        # Sort by position (top to bottom)
        text_elements.sort(key=lambda x: x["y_position"])
        
        # Second pass: categorize and build coherent sections
        for element in text_elements:
            text = element["text"]
            font_size = element["font_size"]
            is_bold = element["is_bold"]
            
            # Determine element type
            if font_size > 16 and not content["title"]:
                content["title"] = text
            elif is_bold and len(text) < 200:  # Likely a heading
                # Save previous section
                if current_section and current_section["content"].strip():
                    content["sections"].append(current_section)
                
                # Start new section
                current_section = {
                    "subtitle": text,
                    "content": ""
                }
            else:
                # Regular text - add to current section or main text
                if current_section:
                    current_section["content"] += text + " "
                else:
                    content["main_text"] += text + " "
        
        # Add final section
        if current_section and current_section["content"].strip():
            content["sections"].append(current_section)
        
        # Clean up whitespace
        content["main_text"] = " ".join(content["main_text"].split())
        for section in content["sections"]:
            section["content"] = " ".join(section["content"].split())
        
        return content
    
    def _analyze_content_type(self, text_blocks: Dict, tables: List, images: List) -> str:
        """Determine the primary content type of the page"""
        if tables and len(tables) > 0:
            if images:
                return "table_and_image"
            else:
                return "table_and_text"
        elif images and len(images) > 0:
            return "text_and_image"
        else:
            return "text"
    
    def _generate_metadata(self, file_path: str, total_pages: int) -> DocumentMetadata:
        """Generate document metadata"""
        import os
        file_size = os.path.getsize(file_path)
        
        return DocumentMetadata(
            filename=os.path.basename(file_path),
            total_pages=total_pages,
            extraction_timestamp=datetime.now().isoformat() + "Z",
            file_size=f"{file_size / (1024*1024):.1f}MB",
            file_type="pdf"
        )
    
    def _generate_summary(self, pages: List[PageContent]) -> Dict[str, Any]:
        """Generate document summary"""
        total_tables = sum(len(page.tables) for page in pages)
        total_images = sum(len(page.images) for page in pages)
        
        # Extract main topics (simplified)
        main_topics = []
        key_metrics = []
        
        for page in pages:
            if page.text_content.get("title"):
                main_topics.append(page.text_content["title"])
            
            # Extract potential metrics from tables
            for table in page.tables:
                for row in table.rows:
                    for key, value in row.items():
                        if any(metric_word in str(value).lower() 
                              for metric_word in ["target", "goal", "%", "reduction"]):
                            key_metrics.append(f"{key}: {value}")
        
        return {
            "total_tables": total_tables,
            "total_images": total_images,
            "main_topics": main_topics[:10],  # Limit to top 10
            "key_metrics": key_metrics[:10]   # Limit to top 10
        }
    
    def _get_file_type(self, file_path: str) -> str:
        """Determine file type from extension"""
        return file_path.split('.')[-1].lower()
    
    def _extract_image(self, file_path: str) -> Dict[str, Any]:
        """Extract text from image files using OCR"""
        try:
            # Load and preprocess image
            img = cv2.imread(file_path)
            if img is None:
                raise ValueError(f"Could not load image: {file_path}")
            
            # Preprocess image for better OCR
            if self.config.enable_preprocessing:
                img = self._preprocess_image(img)
            
            # Convert to PIL Image for tesseract
            img_rgb = cv2.cvtColor(img, cv2.COLOR_BGR2RGB)
            pil_img = Image.fromarray(img_rgb)
            
            # Extract text using OCR
            custom_config = f'--oem 3 --psm 6 -l {self.config.ocr_language}'
            extracted_text = pytesseract.image_to_string(pil_img, config=custom_config)
            
            # Get detailed OCR data for layout analysis
            ocr_data = pytesseract.image_to_data(pil_img, output_type=pytesseract.Output.DICT)
            
            # Analyze layout and extract structured content
            text_content = self._analyze_ocr_layout(ocr_data, extracted_text)
            
            # Try to detect and extract tables
            tables = self._extract_tables_from_image(img, file_path)
            
            # Generate image description
            with open(file_path, 'rb') as f:
                img_data = f.read()
            description = self._generate_image_description(img_data)
            
            # Create image metadata
            height, width = img.shape[:2]
            extracted_image = ExtractedImage(
                image_id="main_image",
                description=description,
                position="full_page",
                estimated_size=f"{width}x{height}",
                file_path=file_path,
                metadata={
                    "width": width,
                    "height": height,
                    "channels": img.shape[2] if len(img.shape) > 2 else 1,
                    "file_size": os.path.getsize(file_path)
                }
            )
            
            # Determine content type
            content_type = "text_and_image" if extracted_text.strip() else "image"
            if tables:
                content_type = "table_and_image"
            
            page_content = PageContent(
                page_number=1,
                content_type=content_type,
                text_content=text_content,
                tables=tables,
                images=[extracted_image]
            )
            
            # Generate metadata
            metadata = self._generate_metadata(file_path, 1)
            metadata.file_type = self._get_file_type(file_path)
            
            # Generate summary
            summary = self._generate_summary([page_content])
            
            return {
                "document_metadata": asdict(metadata),
                "pages": [asdict(page_content)],
                "summary": summary
            }
            
        except Exception as e:
            raise Exception(f"Failed to extract from image {file_path}: {str(e)}")
    
    def _preprocess_image(self, img: np.ndarray) -> np.ndarray:
        """Preprocess image for better OCR results"""
        # Convert to grayscale
        gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
        
        # Remove noise
        denoised = cv2.medianBlur(gray, 3)
        
        # Enhance contrast
        clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8,8))
        enhanced = clahe.apply(denoised)
        
        # Binarization
        _, binary = cv2.threshold(enhanced, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
        
        # Convert back to BGR for consistency
        return cv2.cvtColor(binary, cv2.COLOR_GRAY2BGR)
    
    def _analyze_ocr_layout(self, ocr_data: Dict, full_text: str) -> Dict[str, Any]:
        """Analyze OCR data to extract structured content"""
        content = {
            "title": "",
            "sections": [],
            "main_text": full_text.strip()
        }
        
        # Group text by confidence and size
        high_conf_texts = []
        for i, conf in enumerate(ocr_data['conf']):
            if int(conf) > 60:  # High confidence text
                text = ocr_data['text'][i].strip()
                if text:
                    height = ocr_data['height'][i]
                    high_conf_texts.append((text, height, int(conf)))
        
        # Sort by height (font size proxy) and confidence
        high_conf_texts.sort(key=lambda x: (-x[1], -x[2]))
        
        # Extract potential title (largest, highest confidence text)
        if high_conf_texts:
            content["title"] = high_conf_texts[0][0]
        
        return content
    
    def _extract_tables_from_image(self, img: np.ndarray, file_path: str) -> List[ExtractedTable]:
        """Extract tables from image using computer vision"""
        tables = []
        
        try:
            # Convert to grayscale
            gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
            
            # Apply threshold to get binary image
            _, binary = cv2.threshold(gray, 128, 255, cv2.THRESH_BINARY_INV)
            
            # Detect horizontal lines
            horizontal_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (40, 1))
            horizontal_lines = cv2.morphologyEx(binary, cv2.MORPH_OPEN, horizontal_kernel)
            
            # Detect vertical lines
            vertical_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (1, 40))
            vertical_lines = cv2.morphologyEx(binary, cv2.MORPH_OPEN, vertical_kernel)
            
            # Combine horizontal and vertical lines
            table_mask = cv2.addWeighted(horizontal_lines, 0.5, vertical_lines, 0.5, 0.0)
            
            # Find contours (potential table cells)
            contours, _ = cv2.findContours(table_mask, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
            
            # Filter contours that might represent tables
            table_contours = []
            for contour in contours:
                area = cv2.contourArea(contour)
                if area > 1000:  # Minimum area threshold
                    x, y, w, h = cv2.boundingRect(contour)
                    aspect_ratio = w / h
                    if 0.5 < aspect_ratio < 5:  # Reasonable aspect ratio
                        table_contours.append((x, y, w, h))
            
            # Extract text from detected table regions
            for i, (x, y, w, h) in enumerate(table_contours):
                roi = img[y:y+h, x:x+w]
                roi_text = pytesseract.image_to_string(roi, config='--psm 6')
                
                # Simple table parsing (split by lines and common delimiters)
                lines = [line.strip() for line in roi_text.split('\n') if line.strip()]
                if len(lines) >= self.config.min_table_rows:
                    # Try to detect columns by splitting on common delimiters
                    rows = []
                    headers = []
                    
                    for j, line in enumerate(lines):
                        # Split by multiple spaces, tabs, or |
                        cells = re.split(r'\s{2,}|\t|\|', line)
                        cells = [cell.strip() for cell in cells if cell.strip()]
                        
                        if len(cells) >= self.config.min_table_cols:
                            if j == 0:
                                headers = [f"col_{k}" if not cell else cell for k, cell in enumerate(cells)]
                            else:
                                row_dict = {headers[k] if k < len(headers) else f"col_{k}": cell 
                                          for k, cell in enumerate(cells)}
                                rows.append(row_dict)
                    
                    if rows:
                        table = ExtractedTable(
                            table_id=f"image_table_{i}",
                            headers=headers,
                            rows=rows,
                            confidence=0.7
                        )
                        tables.append(table)
        
        except Exception as e:
            print(f"Table extraction from image failed: {str(e)}")
        
        return tables
    
    def _extract_word(self, file_path: str) -> Dict[str, Any]:
        """Extract content from Word documents"""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required for Word document processing. Install with: pip install python-docx")
        
        try:
            doc = Document(file_path)
            pages = []
            current_page = 1
            
            # Extract paragraphs and structure
            text_content = self._extract_word_text_structure(doc)
            
            # Extract tables
            tables = self._extract_word_tables(doc)
            
            # Extract images (if any)
            images = self._extract_word_images(doc, file_path)
            
            # Determine content type
            content_type = "text"
            if tables and images:
                content_type = "table_and_image"
            elif tables:
                content_type = "table_and_text"
            elif images:
                content_type = "text_and_image"
            
            page_content = PageContent(
                page_number=current_page,
                content_type=content_type,
                text_content=text_content,
                tables=tables,
                images=images
            )
            pages.append(page_content)
            
            # Generate metadata
            metadata = self._generate_metadata(file_path, 1)
            metadata.file_type = "docx"
            
            # Generate summary
            summary = self._generate_summary(pages)
            
            return {
                "document_metadata": asdict(metadata),
                "pages": [asdict(page_content)],
                "summary": summary
            }
            
        except Exception as e:
            raise Exception(f"Failed to extract from Word document {file_path}: {str(e)}")
    
    def _extract_word_text_structure(self, doc) -> Dict[str, Any]:
        """Extract structured text content from Word document"""
        content = {
            "title": "",
            "sections": [],
            "main_text": ""
        }
        
        current_section = None
        all_text = []
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            
            # Check if this is a heading
            if paragraph.style.name.startswith('Heading'):
                if paragraph.style.name == 'Heading 1' and not content["title"]:
                    content["title"] = text
                else:
                    # Save previous section
                    if current_section:
                        content["sections"].append(current_section)
                    
                    # Start new section
                    current_section = {
                        "subtitle": text,
                        "content": ""
                    }
            else:
                # Regular paragraph
                all_text.append(text)
                if current_section:
                    current_section["content"] += " " + text
                else:
                    content["main_text"] += " " + text
        
        # Add last section
        if current_section:
            content["sections"].append(current_section)
        
        # If no title was found in headings, try to extract from first paragraph
        if not content["title"] and all_text:
            content["title"] = all_text[0][:100] + "..." if len(all_text[0]) > 100 else all_text[0]
        
        return content
    
    def _extract_word_tables(self, doc) -> List[ExtractedTable]:
        """Extract tables from Word document"""
        tables = []
        
        for i, table in enumerate(doc.tables):
            try:
                if len(table.rows) < self.config.min_table_rows:
                    continue
                
                # Extract headers from first row
                headers = []
                first_row = table.rows[0]
                for cell in first_row.cells:
                    headers.append(cell.text.strip())
                
                if len(headers) < self.config.min_table_cols:
                    continue
                
                # Extract data rows
                rows = []
                for row in table.rows[1:]:
                    row_data = {}
                    for j, cell in enumerate(row.cells):
                        header = headers[j] if j < len(headers) else f"col_{j}"
                        row_data[header] = cell.text.strip()
                    rows.append(row_data)
                
                if rows:
                    extracted_table = ExtractedTable(
                        table_id=f"word_table_{i}",
                        headers=headers,
                        rows=rows,
                        confidence=0.95
                    )
                    tables.append(extracted_table)
                    
            except Exception as e:
                print(f"Failed to extract table {i}: {str(e)}")
                continue
        
        return tables
    
    def _extract_word_images(self, doc, file_path: str) -> List[ExtractedImage]:
        """Extract images from Word document"""
        images = []
        
        try:
            # Word documents are essentially zip files
            import zipfile
            
            with zipfile.ZipFile(file_path, 'r') as docx_zip:
                # Look for images in the media folder
                image_files = [f for f in docx_zip.namelist() if f.startswith('word/media/')]
                
                for i, img_path in enumerate(image_files):
                    try:
                        # Extract image
                        img_data = docx_zip.read(img_path)
                        
                        # Save image
                        img_filename = f"word_image_{i}.{img_path.split('.')[-1]}"
                        img_output_path = self.images_dir / img_filename
                        
                        if self.config.save_images:
                            with open(img_output_path, 'wb') as f:
                                f.write(img_data)
                        
                        # Generate description
                        description = self._generate_image_description(img_data)
                        
                        extracted_image = ExtractedImage(
                            image_id=f"word_img_{i}",
                            description=description,
                            position="embedded",
                            estimated_size="unknown",
                            file_path=str(img_output_path),
                            metadata={
                                "original_path": img_path,
                                "file_size": len(img_data)
                            }
                        )
                        images.append(extracted_image)
                        
                    except Exception as e:
                        print(f"Failed to extract image {img_path}: {str(e)}")
                        continue
        
        except Exception as e:
            print(f"Failed to extract images from Word document: {str(e)}")
        
        return images
    def _consolidate_page_content(self, page_content: PageContent) -> Dict[str, str]:
        """Create consolidated, coherent content blocks for better RAG performance"""
        consolidated = {}
        
        # 1. Create a full text block
        full_text_parts = []
        
        if page_content.text_content.get("title"):
            full_text_parts.append(f"Title: {page_content.text_content['title']}")
        
        if page_content.text_content.get("main_text"):
            full_text_parts.append(page_content.text_content["main_text"])
        
        for section in page_content.text_content.get("sections", []):
            if section.get("subtitle") and section.get("content"):
                full_text_parts.append(f"{section['subtitle']}: {section['content']}")
        
        consolidated["full_text"] = " ".join(full_text_parts)
        
        # 2. Create structured data summary
        if page_content.tables:
            table_summaries = []
            for table in page_content.tables:
                table_text = f"Table with columns: {', '.join(table.headers)}. "
                
                # Add key data points
                for i, row in enumerate(table.rows[:3]):  # First 3 rows only
                    row_summary = []
                    for key, value in row.items():
                        if value and value.strip():
                            row_summary.append(f"{key}: {value}")
                    
                    if row_summary:
                        table_text += f"Row {i+1}: {'; '.join(row_summary)}. "
                
                if len(table.rows) > 3:
                    table_text += f"... and {len(table.rows) - 3} more rows."
                
                table_summaries.append(table_text)
            
            consolidated["tables_summary"] = " ".join(table_summaries)
        
        # 3. Create combined searchable content
        searchable_parts = [consolidated["full_text"]]
        if consolidated.get("tables_summary"):
            searchable_parts.append(consolidated["tables_summary"])
        
        consolidated["searchable_content"] = " ".join(searchable_parts)
        
        return consolidated


# Usage example and utility functions
def create_extraction_config(
    use_ai: bool = False,
    ai_key: str = None,
    ai_url: str = "https://api.openai.com/v1/chat/completions",
    output_dir: str = "extracted_content",
    table_method: str = "auto"
) -> ExtractionConfig:
    """Create extraction configuration with common settings"""
    return ExtractionConfig(
        use_ai_service=use_ai,
        ai_api_key=ai_key,
        ai_service_url=ai_url,
        output_dir=output_dir,
        table_extraction_method=table_method
    )

def batch_extract_documents(
    file_paths: List[str], 
    config: Optional[ExtractionConfig] = None,
    output_format: str = "json"
) -> Dict[str, Any]:
    """Extract multiple documents in batch"""
    extractor = DocumentExtractor(config)
    results = {}
    
    for file_path in file_paths:
        try:
            print(f"Processing {file_path}...")
            result = extractor.extract_document(file_path)
            results[file_path] = result
            
            # Save individual result
            output_path = extractor.output_dir / f"{Path(file_path).stem}_extracted.json"
            with open(output_path, 'w', encoding='utf-8') as f:
                json.dump(result, f, indent=2, ensure_ascii=False)
                
        except Exception as e:
            print(f"Failed to process {file_path}: {str(e)}")
            results[file_path] = {"error": str(e)}
    
    # Save batch results
    batch_output_path = extractor.output_dir / "batch_results.json"
    with open(batch_output_path, 'w', encoding='utf-8') as f:
        json.dump(results, f, indent=2, ensure_ascii=False)
    
    return results

def main():
    """Example usage of the document extractor"""
    # Basic configuration
    config = create_extraction_config(
        use_ai=False,  # Set to True and provide API key for AI features
        output_dir="extracted_content",
        table_method="auto"
    )
    
    # For AI features, uncomment and configure:
    # config = create_extraction_config(
    #     use_ai=True,
    #     ai_key="your_openai_api_key_here",
    #     ai_url="https://api.openai.com/v1/chat/completions"
    # )
    
    extractor = DocumentExtractor(config)
    
    # Single document extraction
    try:
        file_name = "adm"
        result = extractor.extract_document(f"PDFS/{file_name}.pdf")
        
        output_path = extractor.output_dir/"JSON"/f"{file_name}.json"
        try:
            with open(output_path, "w", encoding='utf-8') as f:
                json.dump(result, f, indent=2, ensure_ascii=False)
            print(f"JSON saved successfully to: {output_path}")
        except Exception as e:
            print(f"Failed to save JSON: {str(e)}")

        print(f"Document extraction completed! Results saved to {output_path}")
        
        # Print summary
        print("\n=== EXTRACTION SUMMARY ===")
        print(f"Total pages: {result['document_metadata']['total_pages']}")
        print(f"Total tables: {result['summary']['total_tables']}")
        print(f"Total images: {result['summary']['total_images']}")
        print(f"Main topics: {', '.join(result['summary']['main_topics'][:3])}")
        
    except Exception as e:
        print(f"Extraction failed: {str(e)}")
    
    # Batch processing example
    # file_list = ["doc1.pdf", "doc2.docx", "image1.png"]
    # batch_results = batch_extract_documents(file_list, config)
    # print(f"Batch processing completed. Processed {len(batch_results)} files.")

if __name__ == "__main__":
    main()
